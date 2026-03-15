"""
File Ingestors
──────────────
Supported formats: .xlsx · .xls · .csv · .pdf · .txt · .md · .rst · .json · .jsonl · .docx

Each ingestor returns a list of Document objects.
The registry maps file extensions → ingestor class (auto-discovery via @register).

Usage:
    docs = load(Path("data.xlsx"))          # auto-selects ExcelIngestor
    docs = load(Path("notes.pdf"))          # auto-selects PDFIngestor

Extend:
    @register(".html")
    class HTMLIngestor:
        def load(self, file_path: Path) -> list[Document]: ...
"""
from __future__ import annotations

import json as _json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


# ─────────────────────────────────────────────────────────────────
# Data model
# ─────────────────────────────────────────────────────────────────

@dataclass
class Document:
    content: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)


# ─────────────────────────────────────────────────────────────────
# Registry
# ─────────────────────────────────────────────────────────────────

_registry: dict[str, type] = {}


def register(*extensions: str):
    """Class decorator — registers an ingestor for one or more file extensions."""
    def decorator(cls):
        for ext in extensions:
            _registry[ext.lower()] = cls
        return cls
    return decorator


def load(file_path: Path) -> list[Document]:
    """Load a file using the appropriate ingestor. Raises ValueError for unsupported types."""
    ext = file_path.suffix.lower()
    cls = _registry.get(ext)
    if cls is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {sorted(_registry.keys())}"
        )
    return cls().load(file_path)


def supported_extensions() -> list[str]:
    return sorted(_registry.keys())


# ─────────────────────────────────────────────────────────────────
# Excel  (.xlsx, .xls)
# ─────────────────────────────────────────────────────────────────

@register(".xlsx", ".xls")
class ExcelIngestor:
    def load(self, file_path: Path) -> list[Document]:
        import pandas as pd

        docs: list[Document] = []
        xl = pd.ExcelFile(file_path)
        for sheet in xl.sheet_names:
            df = xl.parse(sheet).fillna("")
            for idx, row in df.iterrows():
                parts = [f"{col}: {val}" for col, val in row.items() if str(val).strip()]
                if parts:
                    docs.append(
                        Document(
                            content="\n".join(parts),
                            source=str(file_path),
                            metadata={"sheet": sheet, "row": int(idx), "file_type": "excel"},
                        )
                    )
        return docs


# ─────────────────────────────────────────────────────────────────
# CSV  (.csv)
# ─────────────────────────────────────────────────────────────────

@register(".csv")
class CSVIngestor:
    def load(self, file_path: Path) -> list[Document]:
        import pandas as pd

        df = pd.read_csv(file_path).fillna("")
        docs: list[Document] = []
        for idx, row in df.iterrows():
            parts = [f"{col}: {val}" for col, val in row.items() if str(val).strip()]
            if parts:
                docs.append(
                    Document(
                        content="\n".join(parts),
                        source=str(file_path),
                        metadata={"row": int(idx), "file_type": "csv"},
                    )
                )
        return docs


# ─────────────────────────────────────────────────────────────────
# PDF  (.pdf)
# ─────────────────────────────────────────────────────────────────

@register(".pdf")
class PDFIngestor:
    def load(self, file_path: Path) -> list[Document]:
        import fitz  # pymupdf

        docs: list[Document] = []
        with fitz.open(str(file_path)) as pdf:
            for page_num in range(len(pdf)):
                text = pdf[page_num].get_text()
                if text.strip():
                    docs.append(
                        Document(
                            content=text,
                            source=str(file_path),
                            metadata={"page": page_num + 1, "file_type": "pdf"},
                        )
                    )
        return docs


# ─────────────────────────────────────────────────────────────────
# Plain text / Markdown / RST  (.txt, .md, .rst)
# ─────────────────────────────────────────────────────────────────

@register(".txt", ".md", ".rst")
class TextIngestor:
    def load(self, file_path: Path) -> list[Document]:
        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        paragraphs = [p.strip() for p in raw.split("\n\n") if p.strip()]
        return [
            Document(
                content=para,
                source=str(file_path),
                metadata={"paragraph": i, "file_type": "text"},
            )
            for i, para in enumerate(paragraphs)
        ]


# ─────────────────────────────────────────────────────────────────
# JSON / JSONL  (.json, .jsonl)
# ─────────────────────────────────────────────────────────────────

@register(".json", ".jsonl")
class JSONIngestor:
    def load(self, file_path: Path) -> list[Document]:
        raw = file_path.read_text(encoding="utf-8")
        items: list[Any]

        if file_path.suffix == ".jsonl":
            items = [_json.loads(line) for line in raw.splitlines() if line.strip()]
        else:
            data = _json.loads(raw)
            items = data if isinstance(data, list) else [data]

        return [
            Document(
                content=_json.dumps(item, indent=2, ensure_ascii=False),
                source=str(file_path),
                metadata={"item_index": i, "file_type": "json"},
            )
            for i, item in enumerate(items)
        ]


# ─────────────────────────────────────────────────────────────────
# Word Document  (.docx)
# ─────────────────────────────────────────────────────────────────

@register(".docx")
class DocxIngestor:
    def load(self, file_path: Path) -> list[Document]:
        from docx import Document as DocxDoc

        word = DocxDoc(str(file_path))
        docs: list[Document] = []
        section_lines: list[str] = []
        section_idx = 0

        for para in word.paragraphs:
            text = para.text.strip()
            if text:
                section_lines.append(text)
            elif section_lines:
                docs.append(
                    Document(
                        content="\n".join(section_lines),
                        source=str(file_path),
                        metadata={"section": section_idx, "file_type": "docx"},
                    )
                )
                section_lines = []
                section_idx += 1

        if section_lines:
            docs.append(
                Document(
                    content="\n".join(section_lines),
                    source=str(file_path),
                    metadata={"section": section_idx, "file_type": "docx"},
                )
            )

        return docs
